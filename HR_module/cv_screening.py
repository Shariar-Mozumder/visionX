from tempfile import NamedTemporaryFile
from typing import List
import PyPDF2
from fastapi.responses import JSONResponse
import spacy
from sentence_transformers import SentenceTransformer, util
import re
from datetime import datetime
import os
import zipfile
from bs4 import BeautifulSoup
from docx import Document

from fastapi import FastAPI, File, Path, UploadFile
import uvicorn
app = FastAPI()

import sys
import os

# sys.path.append(os.path.abspath("agents"))

from HR_Module_agents import  process_resume,score_resume

# def extract_text_from_pdf(pdf_path):
#     with open(pdf_path, 'rb') as file:
#         reader = PyPDF2.PdfReader(file)
#         text = " ".join([page.extract_text() for page in reader.pages])
#     return text


import pdfplumber

def extract_text_from_file(file_path):
    file_extension = os.path.splitext(file_path)[1].lower()
    extracted_text = ""

    try:
        if file_extension == '.pdf':
            # Extract text from PDF
            # with open(file_path, 'rb') as file:
            #     reader = PyPDF2.PdfReader(file)
            #     extracted_text = " ".join([page.extract_text() for page in reader.pages])

            with pdfplumber.open(file_path) as pdf:
                extracted_text = ""
                for page in pdf.pages:
                    # Extract text with layout information
                    extracted_text += page.extract_text()

        elif file_extension in ['.docx', '.doc']:
            # Extract text from Word document
            doc = Document(file_path)
            # extracted_text = " ".join([paragraph.text for paragraph in doc.paragraphs])
            paragraphs = [paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()]
            tables = []
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        tables.append(cell.text.strip())
            # Combine all extracted text
            extracted_text = " ".join(paragraphs + tables)

        elif file_extension == '.txt':
            # Extract text from plain text file
            with open(file_path, 'r', encoding='utf-8') as file:
                extracted_text = file.read()

        elif file_extension == '.html':
            # Extract text from HTML file
            with open(file_path, 'r', encoding='utf-8') as file:
                soup = BeautifulSoup(file, 'html.parser')
                extracted_text = soup.get_text()

        elif file_extension == '.zip':
            # Extract text from ZIP file
            with zipfile.ZipFile(file_path, 'r') as zip_file:
                for name in zip_file.namelist():
                    if name.endswith(('.pdf', '.docx', '.doc', '.txt', '.html')):
                        with zip_file.open(name) as file:
                            nested_extension = os.path.splitext(name)[1].lower()
                            nested_text = ""
                            if nested_extension == '.pdf':
                                reader = PyPDF2.PdfReader(file)
                                nested_text = " ".join([page.extract_text() for page in reader.pages])
                            elif nested_extension in ['.docx', '.doc']:
                                doc = Document(file)
                                nested_text = " ".join([paragraph.text for paragraph in doc.paragraphs])
                            elif nested_extension == '.txt':
                                nested_text = file.read().decode('utf-8')
                            elif nested_extension == '.html':
                                soup = BeautifulSoup(file, 'html.parser')
                                nested_text = soup.get_text()
                            extracted_text += nested_text + "\n"

        else:
            raise ValueError(f"Unsupported file type: {file_extension}")

    except Exception as e:
        return f"Error while processing {file_path}: {e}"

    return extracted_text




nlp = spacy.load("en_core_web_sm")

def preprocess_text(text):
    doc = nlp(text.lower())
    return " ".join([token.lemma_ for token in doc if not token.is_stop and token.is_alpha])



model = SentenceTransformer('sentence-transformers/all-MiniLM-L6-v2')

def compute_similarity(job_description, resume_text):
    job_embedding = model.encode(job_description, convert_to_tensor=True)
    resume_embedding = model.encode(resume_text, convert_to_tensor=True)
    return util.pytorch_cos_sim(job_embedding, resume_embedding).item()



def extract_skills(text):
    skills_list = ["python", "machine learning", "data analysis", "communication"]  # Add more skills
    return [skill for skill in skills_list if skill in text.lower()]

# def extract_entities(resume_text):
#     doc = nlp(resume_text)
#     entities = {"Skills": [], "Education": [], "Experience": []}
#     for ent in doc.ents:
#         if ent.label_ in ["ORG", "EDUCATION"]:
#             entities["Education"].append(ent.text)
#         elif ent.label_ in ["DATE"]:
#             entities["Experience"].append(ent.text)
#         elif ent.label_ in ["PRODUCT", "WORK_OF_ART"]:
#             entities["Skills"].append(ent.text)
#     return entities
def extract_entities(resume_text):
    doc = nlp(resume_text)
    entities = {
        "Work Experience": [],
        "Projects": [],
        "Research": [],
        "Academic Achievements": [],
        "Education": [],
        "Skills": [],
        "Social Work/Volunteer": [],
    }

    for ent in doc.ents:
        # Classify entities based on their labels
        if ent.label_ in ["ORG", "EDUCATION"]:
            entities["Education"].append(ent.text)
        elif ent.label_ == "DATE":
            entities["Work Experience"].append(ent.text)
        elif ent.label_ in ["PRODUCT", "WORK_OF_ART"]:
            entities["Skills"].append(ent.text)
        elif ent.label_ in ["GPE", "LOC"] and "project" in ent.text.lower():
            entities["Projects"].append(ent.text)
        elif ent.label_ in ["PERSON", "ORG"] and "research" in ent.text.lower():
            entities["Research"].append(ent.text)
        elif ent.label_ == "MONEY" or "achievement" in ent.text.lower():
            entities["Academic Achievements"].append(ent.text)
        elif ent.label_ == "NORP" or "volunteer" in ent.text.lower():
            entities["Social Work/Volunteer"].append(ent.text)

    # Heuristic-based enhancements
    for line in resume_text.splitlines():
        line_lower = line.lower()
        if "intern" in line_lower or "work experience" in line_lower:
            entities["Work Experience"].append(line.strip())
        if "project" in line_lower:
            entities["Projects"].append(line.strip())
        if "research" in line_lower:
            entities["Research"].append(line.strip())
        if "achievement" in line_lower or "award" in line_lower:
            entities["Academic Achievements"].append(line.strip())
        if "volunteer" in line_lower or "social work" in line_lower:
            entities["Social Work/Volunteer"].append(line.strip())
        if "skill" in line_lower or "proficient in" in line_lower:
            entities["Skills"].append(line.strip())

    # Deduplicate and clean up extracted entities
    for category in entities:
        entities[category] = list(set(entities[category]))

    return entities


def weighted_keyword_match(resume_text, keywords):
    score = 0
    for keyword, weight in keywords.items():
        if keyword.lower() in resume_text.lower():
            score += weight
    return score

def calculate_experience(resume_text):
    dates = re.findall(r"\b(19|20)\d{2}\b", resume_text)
    if len(dates) >= 2:
        years = [int(date) for date in dates]
        experience = max(years) - min(years)
        return experience
    return 0


def extract_resumes(resumes):
    results = []
    for resume_text in resumes:
        
        # preprocessed_text = preprocess_text(resume_text)
        # similarity_score = compute_similarity(job_description, preprocessed_text)
        # entities=extract_entities(resume_text)
        # exp=calculate_experience(resume_text)
        ai_rsponse=process_resume(resume_text)
        print(ai_rsponse)
        results.append(ai_rsponse)
        # results.append((resume_path, similarity_score))
    return results

def resumes_scores_count(job_description, resume_list):
    scores=[]
    for resume in resume_list:
        score_result=score_resume(resume,job_description)
        # Contact_Details=resume.get('Contact Details')
        resume_sores={
            # "Contact_Details":Contact_Details,
            "score_result": score_result
        }
        scores.append(resume_sores)
    return scores

def rank_resume(job_description,resumes_text):
    # needs to elaborate more
    resume_list=extract_resumes(resumes_text)
    score_result=resumes_scores_count(job_description,resume_list)
    return resume_list,score_result



def calculate_total_score(similarity, keyword_score, experience, soft_skills_score):
    return (
        0.4 * similarity +
        0.3 * keyword_score +
        0.2 * experience +
        0.1 * soft_skills_score
    )

# Example inputs
similarity = 0.85  # Similarity score from Sentence Transformers
keyword_score = 25
experience = 5  # Years
soft_skills_score = 0.9  # Sentiment analysis score

total_score = calculate_total_score(similarity, keyword_score, experience, soft_skills_score)



# if __name__=="__main__":
#     resumes_path=[
#             'C:/Users/Lenovo/Downloads/Md_Shariar_Hossain_ML_CV.pdf',
#              'C:/Users/Lenovo/Downloads/Samin-Yasar-Chowdhury_CV.pdf',
#              'C:/Users/Lenovo/Downloads/Farjana_cv.pdf',
#              'C:/Users/Lenovo/Downloads/doc_cv.docx',
#              'C:/Users/Lenovo/Downloads/cv.txt']
    
#     job_des_path='Job_description.txt'
#     Job_des_extracted_text=''
#     with open(job_des_path, 'r', encoding='utf-8') as file:
#                 Job_des_extracted_text = file.read()
    
#     resumes=[]
#     similarity_scores=[]
#     ranking_report=[]
#     for path in resumes_path:
#         resume_text = extract_text_from_file(path)
#         similarities=compute_similarity(Job_des_extracted_text,resume_text)
#         similarity_scores.append(similarities)
#         resumes.append(resume_text)

#     resume_list,score_result=rank_resume(Job_des_extracted_text,resumes)
#     for i in range(len(resumes_path)):
#         analysis_score_report={
#             "CV_Path":resumes_path[i],
#             "resume_list":resume_list[i],
#             "similarity_scores":similarity_scores[i],
#             "Analysis score_result":score_result[i]
#         }
#         ranking_report.append(analysis_score_report)

#     print(ranking_report)



# @app.post("/resume_screening/")
# async def resume_screening(
#     job_description: UploadFile = File(...), 
#     resumes: List[UploadFile] = File(...)
# ):
#     try:
#         # Read and save the job description
#         job_desc_content = await job_description.read()
#         Job_des_extracted_text = job_desc_content.decode("utf-8")

#         resumes_texts = []
#         similarity_scores = []
#         ranking_report = []

#         # Process each uploaded resume
#         for resume_file in resumes:
#             # Save file temporarily for processing
#             with NamedTemporaryFile(delete=False, suffix=Path(resume_file.filename).suffix) as temp_file:
#                 temp_file.write(await resume_file.read())
#                 temp_file_path = temp_file.name

#             # Extract text from the resume file
#             resume_text = extract_text_from_file(temp_file_path)
#             resumes_texts.append(resume_text)

#             # Compute similarity score
#             similarity = compute_similarity(Job_des_extracted_text, resume_text)
#             similarity_scores.append(similarity)

#             # Clean up the temporary file
#             os.remove(temp_file_path)

#         # Rank resumes based on similarity
#         resume_list, score_result = rank_resume(Job_des_extracted_text, resumes_texts)

#         # Prepare the ranking report
#         for i in range(len(resumes)):
#             analysis_score_report = {
#                 "CV_Name": resumes[i].filename,
#                 "resume_list": resume_list[i],
#                 "similarity_score": similarity_scores[i],
#                 "analysis_score_result": score_result[i]
#             }
#             ranking_report.append(analysis_score_report)

#         return JSONResponse(content={"ranking_report": ranking_report})

#     except Exception as e:
#         return JSONResponse(status_code=500, content={"error": str(e)})

# if __name__ == "__main__":
#     uvicorn.run(app, host="0.0.0.0", port=8000)
from sqlalchemy import create_engine, text
DATABASE_URL = 'mssql+pyodbc://DESKTOP-2419RQF/visionX?driver=ODBC+Driver+17+for+SQL+Server&trusted_connection=yes'
engine = create_engine(DATABASE_URL)
from db_operations import insert_candidate


@app.post("/resume_screening/")
async def resume_screening(
    job_description: UploadFile = File(...), 
    resumes: List[UploadFile] = File(...)
):
    try:
        # Read the job description content (expecting a .txt file)
        job_desc_content = await job_description.read()
        Job_des_extracted_text = job_desc_content.decode("utf-8")

        resumes_texts = []
        similarity_scores = []
        ranking_report = []

        # Process each uploaded resume
        for resume_file in resumes:
            # Determine file extension from the uploaded file
            file_extension = os.path.splitext(resume_file.filename)[1]

            # Save file temporarily for processing with the correct extension
            with NamedTemporaryFile(delete=False, suffix=file_extension) as temp_file:
                temp_file.write(await resume_file.read())
                temp_file_path = temp_file.name

            # Extract text from the resume file
            resume_text = extract_text_from_file(temp_file_path)
            resumes_texts.append(resume_text)

            # Compute similarity score
            similarity = compute_similarity(Job_des_extracted_text, resume_text)
            similarity_scores.append(similarity)

            # Clean up the temporary file
            os.remove(temp_file_path)

        # Rank resumes based on similarity
        resume_list, score_result = rank_resume(Job_des_extracted_text, resumes_texts)

        # Prepare the ranking report
        for i in range(len(resumes)):
            candidate_data = {
                "ContactInformation": str(resume_list[i][0].get("Contact Details")),
                "AcademicEducation": str(resume_list[i][0].get("Education")),
                "WorkExperience": str(resume_list[i][0].get("Work Experience")),
                "Skills": str(resume_list[i][0].get("Skills and Certifications")),
                "CompatibilityScore": score_result[i].get("score_result")[0].get("Compatibility_Score"),
                "SimilaritSscore": similarity_scores[i],
                "Stage": "screening",
                "Shortlisted": "Not decided",
                "ScreeningDate": "1900-01-01T10:30:00",
                "ScreeningResult": "Not decided",
                "TechnicalTestDate": "1900-01-01T10:30:00",
                "TechnicalTestResult": "Not decided",
                "HrTestDate": "1900-01-01T10:30:00",
                "HrTestResult": "Not decided",
                "JoiningDate": "1900-01-01T10:30:00",
                "Joined": "No",
                "Blacklisted": "No",
                "JobDescription": Job_des_extracted_text,
            }

            res=insert_candidate(candidate_data)
            analysis_score_report = {
                "CV_Name": resumes[i].filename,
                "resume_list": resume_list[i],
                "similarity_score": similarity_scores[i],
                "analysis_score_result": score_result[i]
            }
            ranking_report.append(analysis_score_report)

        return JSONResponse(content={"ranking_report": ranking_report})

    except Exception as e:
        return JSONResponse(status_code=500, content={"error": str(e)})

if __name__ == "__main__":
    uvicorn.run(app, host="0.0.0.0", port=8000)